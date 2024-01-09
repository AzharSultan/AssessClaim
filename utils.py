import re
import docx
import backoff
import openai
import numpy as np
import pandas as pd
import logging

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

def get_record_df(record_path):
    doc = docx.Document(record_path)
    df = pd.DataFrame()
    df['text'] = [p.text for p in doc.paragraphs]
    df['type'] = [p.style for p in doc.paragraphs]
    return df

def extract_n_digit_number(sentence, n=4):
    pattern = fr'\b\d{{{n}}}\b'
    matches = re.findall(pattern, sentence)
    return matches

def log_retry(x):
    logger.info('retrying')
    return x

@backoff.on_exception(backoff.expo,
                      (openai.APIError,
                       openai.APIConnectionError,
                       openai.ConflictError,
                       openai.RateLimitError,
                       openai.InternalServerError,
                       openai.APITimeoutError),
                      on_backoff=log_retry,
                      max_value=30,
                      max_tries=15,
                      jitter=None)
def chat_with_gpt(user_prompt, system_prompt=None, model='gpt-4'):
    if system_prompt is None:
        system_prompt = "You are a helpful assistant."
    client = openai.OpenAI(timeout=60)
    response = client.chat.completions.create(
        model=model,
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt}
        ],
        temperature=0,
    )
    logger.info(user_prompt)
    logger.info(f'LLM response: {response.choices[0].message.content}')

    return response.choices[0].message.content

def get_initial_answers(medical_history):
    pre_prompt = 'Based on the patient information provided below, answer the questions at the end: \n'
    post_prompt = """Based on the information above, answer the questions below in order:
1. What year was the patient born in?
2. What is the 5-digit code for requested procedure?
3. Does the patient have rectal bleeding?
4. Does the patient have iron deficiency anemia?
5. Does the patient have abdominal pain?
6. Does the patient have Telangiectasia?
7. Has the patient had colonoscopy before?
8. Did the patient receive any treatment for his current condition?
9. Did anyone in the patient's family have colorectal cancer?
10. Did anyone in the patient's family have colonic adenomatous polyposis of unknown etiology?

Format for the answer should be 'Yes/No/Not possible to determine, reason for the answer'"""
    prompt = '\n\n'.join([pre_prompt, medical_history, post_prompt])
    response = chat_with_gpt(prompt)
    answers = response.split('\n')
    if len(answers) != 10:
        # TODO: More sophisticated output parsing
        raise NotImplementedError('Answer parsing from LLM needs to be improved')
    return answers

def get_verification_statements(medical_history, birth_year, cpt=45378):
    pre_prompt = 'Based on the patient information provided below, answer the questions at the end: \n'
    
    post_prompt = (f"Based on the information above, identify whether each of the statement below is true or false:\n",
    f"1. The patient was born in year {birth_year}.",
    f"2. {cpt} is not one of the requested procedures.",
    f"3. The patient has had a colonoscopy in the last 10 years.",
    f"4. The patient recently received a treatment (including over the counter drugs) for his condition.",
    f"5. The patient has one of the these symptoms: abdominal pain, rectal bleeding, iron deficiency anemia or Telangiectasia",
    f"6. The patient does not have a family history of colorectal cancer.",
    f"7. The patient recently received a treatment for his condition and it improved the symptoms.",
    f"8. Someone in the patient's immediate family had colonic adenomatous polyposis of unknown etiology",
    f"9. The patient recently received a treatment for his condition and it did not improve the symptoms."
    f"\nOnly answer true or false for each statement, do not explain the answer")
    post_prompt = '\n'.join(post_prompt)
    prompt = '\n\n'.join([pre_prompt, medical_history, post_prompt])
    response = chat_with_gpt(prompt)
    answers = response.split('\n')
    if len(answers) != 9:
        # TODO: More sophisticated output parsing
        raise NotImplementedError('Answer parsing from LLM needs to be improved')
    return answers

def extract_age(initial_answer, statement, rag_model, match_threshold=16, current_year=2023):
    year = extract_n_digit_number(initial_answer, 4)
    dob = int(year[0])
    age = current_year - dob
    search_match = rag_model.search(f'Date of Birth DOB {dob}')[0]
    is_age_search = (search_match['score'] > match_threshold) & (str(dob) in search_match['content'])
    evidence = search_match['content']
    
    if 'true' in statement.lower():
        conf = int(is_age_search)
    else:
        conf = -1
    
    return age, conf, evidence

def extract_cpt(initial_answer, statement, rag_model, cpt=45378, match_threshold=16):
    codes = extract_n_digit_number(initial_answer, len(str(cpt)))
    codes = [int(c) for c in codes]
    is_cpt_llm = cpt in codes
    if not is_cpt_llm:
        # verify again cpt not present with gpt4
        pass
    else:
        # verify again cpt present with gpt4
        pass
    search_match = rag_model.search(f'Requested procedure: {cpt}')[0]
    is_cpt_search = (search_match['score'] > match_threshold) & (str(cpt) in search_match['content'])
    is_cpt = (is_cpt_llm or is_cpt_search)
    evidence = search_match['content']
    
    if is_cpt != ('false' in statement.lower()):
        conf = -1
    else:
        conf = int(is_cpt_llm == is_cpt_search)

    return is_cpt, conf, evidence

def ask_treatment_outcome(medical_history):
    pre_prompt = "Based on the patient notes provided below, answer the questions at the end:"
    post_prompt = "Based on the information above, answer the questions below in order:\n"\
                  "1. Has the condition of the patient improved based on recent treatment? Disregard any planned treatments\n\n"\
                  "Format for the answer should be 'Yes/No/Not possible to determine, reason for the answer'"
    prompt = '\n\n'.join([pre_prompt, medical_history, post_prompt])
    response = chat_with_gpt(prompt)
    return response

def extract_cons_treatment(initial_answer, statements, medical_history, rag_model, match_threshold=16):
    conf = 0
    evidence = ''
    
    if 'yes, ' not in initial_answer.lower():
        # TODO: Double check that treatment is not mentioned in the report
        reason = 'No evidence of conservative treatment found in the medical record'
        if 'true' in statements[0].lower():
            conf = -1
        is_continue = False
        return is_continue, reason, conf, evidence
    
    evidence = initial_answer.lower().split('yes, ')[-1]
    search_match = rag_model.search(evidence)[0]
    conf = int(search_match['score'] > match_threshold)
    evidence = search_match['content']
    answer = ask_treatment_outcome(medical_history)
    if 'yes, ' in answer.lower():
        reason = 'Patient had a conservative treatment that improved condition as mentioned in the report here:\n'\
                f'{evidence}'
        is_continue = False
        if ('false' in statements[1].lower()) or ('true' in statements[2].lower()):
            conf = -1 
    else:
        reason = None
        is_continue = True
        if ('true' in statements[1].lower()) or ('false' in statements[2].lower()):
            conf = -1
    return is_continue, reason, conf, evidence

def extract_polyposis(initial_answer, statement, rag_model, match_threshold=16):
    is_polyposis = 'yes, ' in initial_answer.lower()
    search_match = rag_model.search(f'family colonic adenomatous polyposis of unknown etiology')[0]
    if search_match['score'] > match_threshold:
        #TODO: check again for polyposis with llm
        evidence = search_match['content']
    else:
        #TODO: check again for absence of polyposis with llm
        #TODO: get evidence from llm
        evidence = search_match['content']
    if is_polyposis != ('true' in statement.lower()):
        conf = -1
    else:
        conf = int(is_polyposis == (search_match['score'] > match_threshold))
    return is_polyposis, conf, evidence

def extract_symptomatic(initial_answers, statement, rag_model, match_threshold=16):
    conf = 0
    evidence = []
    
    is_bleeding = 'yes, ' in initial_answers[0].lower()
    if is_bleeding:
        search_match = rag_model.search(f'Symptoms rectal bleeding')[0]
        if search_match['score'] > match_threshold:
            conf += 1
            evidence.append(search_match['content'])
    
    is_anemia = 'yes, ' in initial_answers[1].lower()
    if is_anemia:
        search_match = rag_model.search(f'Symptoms iron deficiency anemia')[0]
        if search_match['score'] > match_threshold:
            conf += 1
            evidence.append(search_match['content'])
    
    is_pain = 'yes, ' in initial_answers[2].lower()
    if is_pain:
        search_match = rag_model.search(f'Symptoms abdominal discomfort pain')[0]
        if search_match['score'] > match_threshold:
            conf += 1
            evidence.append(search_match['content'])
    
    is_talen = 'yes, ' in initial_answers[3].lower()
    if is_pain:
        search_match = rag_model.search(f'Symptoms Telangiectasia')[0]
        if search_match['score'] > match_threshold:
            conf += 1
            evidence.append(search_match['content'])
    
    is_symptomatic = is_bleeding or is_anemia or is_pain or is_talen
    evidence = '\n'.join(evidence)
    if is_symptomatic != ('true' in statement.lower()):
        conf = -1
    
    return is_symptomatic, conf, evidence

def extract_prior_colonoscopy(initial_answer, statement, rag_model, match_threshold=16):
    conf = 0
    evidence = ''
    is_colon = 'yes, ' in initial_answer.lower()
    if is_colon:
        evidence = initial_answer.lower().split('yes, ')[-1]
        search_match = rag_model.search(evidence)[0]
        conf = int(search_match['score'] > match_threshold)
        evidence = search_match['content']
    
    if is_colon != ('true' in statement.lower()):
        conf = -1
        
    return is_colon, conf, evidence

def extract_cancer_history(initial_answer, statement, rag_model, match_threshold=16):
    conf = 0
    evidence = ''
    is_cancer = 'yes, ' in initial_answer.lower()
    if is_cancer:
        evidence = initial_answer.lower().split('\n')[-1]
        search_match = rag_model.search(evidence)[0]
        conf = int(search_match['score'] > match_threshold)
        evidence = search_match['content']
    
    if is_cancer != ('false' in statement.lower()):
        conf = -1

    return is_cancer, conf, evidence
    